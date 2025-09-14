import random

# Modified: 2025-09-11 19:31:57
"""
程序概述
========
这是一个专为小学生设计的数学练习题自动生成器。它能够创建大数（千位至百万位）的
加、减、乘、除及混合运算题目，并生成格式精美、可直接打印的Word文档。

核心功能
========
1. 题目生成
   - **运算类型**: 支持加法、减法、乘法、除法以及两步混合运算。
   - **难度分级**:
     - `Easy` (简单): 千位数的基本运算。
     - `Medium` (中等): 万位数的基本运算。
     - `Hard` (困难): 十万位数的基本运算。
     - `Harder` (更难): 百万位数的混合运算。
   - **智能分页**: 自动调整题目总数，避免最后一页题目过少，节约纸张。

2. Word文档输出
   - **专业排版**: 题目文档采用双栏布局，答案文档采用四栏布局，优化空间利用率。
   - **题目与答案分离**: 生成两个独立的Word文件，一份用于学生练习，一份用于家长/教师校对。
   - **精美页眉页脚**:
     - 页眉包含标题、难度星级（*至****）和日期占位符。
     - 页脚包含生成时间、页码（当前页/总页数）和难度图例。
   - **格式清晰**:
     - 题号采用小数点对齐，视觉上整洁美观。
     - 通过不同灰度的“=”符号或答案颜色来区分难度，一目了然。

3. 高度可定制
   - **命令行接口**: 提供丰富的命令行参数，可轻松定制题目数量、难度和输出路径。
   - **灵活的难度选择**:
     - 可选择一个或多个难度级别混合出题。
     - 支持 `random` 模式，随机混合所有难度。
     - 支持数字别名（1-5）快速指定难度，简化输入。

使用方法
========
通过命令行运行脚本，并使用参数进行配置。

# 示例 1: 生成20道简单题目
python bigNumber_prolems.py -c 20 -d easy -o ./output

# 示例 2: 生成30道简单和困难混合题目
python bigNumber_prolems.py -c 30 -d easy hard -o ./output

# 示例 3: 使用数字别名生成40道中等和更难的题目
python bigNumber_prolems.py -c 40 -d 2 4 -o ./output

# 示例 4: 生成25道随机难度的题目
python bigNumber_prolems.py -c 25 -d random -o ./output
"""

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor  # 添加RGBColor导入
except ImportError:
    print("The 'python-docx' library is required to run this script.")
    print("Please install it using\n pip install python-docx")
    exit(1)
import argparse
import math
import os
import random
import sys
import tempfile
from datetime import datetime
from os import path


class LargeNumberMathGenerator:
    def __init__(self):
        self.problems = []
        self.answers = []
        self.selected_difficulty = []
        self.available_fonts = [
            'Calibri',
            'Times New Roman',
            'Arial',
            'Consolas',
            'Cambria',  # 常规字体
            'Segoe Print',
            'Comic Sans MS',
            'Segoe Script',
            'Bradley Hand ITC',
            'Ink Free',  # 手写字体
        ]
        self.difficulty_colors = {
            "easy": RGBColor(192, 192, 192),  # 浅灰色 - 简单
            "medium": RGBColor(128, 128, 128),  # 中灰色 - 中等
            "hard": RGBColor(64, 64, 64),  # 深灰色 - 困难
            "harder": RGBColor(0, 0, 0),  # 黑色 - 更难（四则混合）
        }
        self.difficulty_names = {"easy": "简单", "medium": "中等", "hard": "困难", "harder": "更难（四则混合）"}

    def generate_addition(self, difficulty):
        """生成加法题目"""
        if difficulty == "easy":
            num1 = random.randint(1000, 9999)
            num2 = random.randint(1000, 9999)
        elif difficulty == "medium":
            num1 = random.randint(10000, 99999)
            num2 = random.randint(10000, 99999)
        elif difficulty == "hard":
            num1 = random.randint(100000, 999999)
            num2 = random.randint(100000, 999999)
        else:  # harder
            num1 = random.randint(1000000, 9999999)
            num2 = random.randint(1000000, 9999999)

        problem = f"{num1} + {num2} "
        answer = num1 + num2
        return problem, answer, difficulty

    def generate_subtraction(self, difficulty):
        """生成减法题目"""
        if difficulty == "easy":
            num1 = random.randint(5000, 9999)
            num2 = random.randint(1000, num1 - 1000)
        elif difficulty == "medium":
            num1 = random.randint(50000, 99999)
            num2 = random.randint(10000, num1 - 10000)
        elif difficulty == "hard":
            num1 = random.randint(500000, 999999)
            num2 = random.randint(100000, num1 - 100000)
        else:  # harder
            num1 = random.randint(5000000, 9999999)
            num2 = random.randint(1000000, num1 - 1000000)

        problem = f"{num1} - {num2} "
        answer = num1 - num2
        return problem, answer, difficulty

    def generate_multiplication(self, difficulty):
        """生成乘法题目"""
        if difficulty == "easy":
            num1 = random.randint(100, 999)
            num2 = random.randint(10, 99)
        elif difficulty == "medium":
            num1 = random.randint(1000, 9999)
            num2 = random.randint(100, 999)
        elif difficulty == "hard":
            num1 = random.randint(10000, 99999)
            num2 = random.randint(1000, 9999)
        else:  # harder
            num1 = random.randint(100000, 999999)
            num2 = random.randint(10000, 99999)

        problem = f"{num1} × {num2} "
        answer = num1 * num2
        return problem, answer, difficulty

    def generate_division(self, difficulty):
        """生成除法题目"""
        if difficulty == "easy":
            divisor = random.randint(2, 9)
            quotient = random.randint(100, 999)
            dividend = divisor * quotient
        elif difficulty == "medium":
            divisor = random.randint(10, 99)
            quotient = random.randint(100, 999)
            dividend = divisor * quotient
        elif difficulty == "hard":
            divisor = random.randint(100, 999)
            quotient = random.randint(100, 999)
            dividend = divisor * quotient
        else:  # harder
            divisor = random.randint(1000, 9999)
            quotient = random.randint(100, 999)
            dividend = divisor * quotient

        problem = f"{dividend} ÷ {divisor} "
        answer = quotient
        return problem, answer, difficulty

    def generate_mixed_operation(self, difficulty):
        """生成四则混合运算题目"""
        if difficulty == "harder":
            # 生成两种运算的混合题目
            operations = ['+', '-', '×', '÷']
            op1 = random.choice(operations)
            op2 = random.choice(operations)

            max_attempts = 10  # 最大尝试次数，避免无限循环
            attempt = 0

            while attempt < max_attempts:
                attempt += 1

                if op1 in ['×', '÷'] or op2 in ['×', '÷']:
                    # 包含乘除法的混合运算
                    if random.choice([True, False]):
                        # 形式: a × b + c 或 a ÷ b + c
                        if op1 == '×':
                            num1 = random.randint(100, 999)
                            num2 = random.randint(10, 99)
                            part1 = num1 * num2
                        else:  # 除法
                            divisor = random.randint(2, 9)
                            quotient = random.randint(100, 999)
                            num1 = divisor * quotient
                            num2 = divisor
                            part1 = quotient

                        num3 = random.randint(1000, 9999)
                        if op2 == '+':
                            answer = part1 + num3
                            problem = f"{num1} {op1} {num2} {op2} {num3} "
                            return problem, answer, difficulty
                        else:  # 确保减法结果为正数且有合理的范围
                            if part1 > 1000:
                                num3 = random.randint(1000, part1 - 100)
                                answer = part1 - num3
                                problem = f"{num1} {op1} {num2} {op2} {num3} "
                                return problem, answer, difficulty
                    else:
                        # 形式: a + b × c 或 a + b ÷ c
                        num1 = random.randint(1000, 9999)
                        if op2 == '×':
                            num2 = random.randint(10, 99)
                            num3 = random.randint(10, 99)
                            part2 = num2 * num3
                        else:  # 除法
                            divisor = random.randint(2, 9)
                            quotient = random.randint(100, 999)
                            num2 = divisor * quotient
                            num3 = divisor
                            part2 = quotient

                        if op1 == '+':
                            answer = num1 + part2
                            problem = f"{num1} {op1} {num2} {op2} {num3} "
                            return problem, answer, difficulty
                        else:  # 确保减法结果为正数
                            if num1 > part2:
                                answer = num1 - part2
                                problem = f"{num1} {op1} {num2} {op2} {num3} "
                                return problem, answer, difficulty
                else:
                    # 加减法混合运算
                    num1 = random.randint(1000, 9999)
                    num2 = random.randint(1000, 9999)
                    num3 = random.randint(1000, 9999)

                    if op1 == '+' and op2 == '+':
                        answer = num1 + num2 + num3
                        problem = f"{num1} {op1} {num2} {op2} {num3} "
                        return problem, answer, difficulty
                    elif op1 == '+' and op2 == '-':
                        if num1 + num2 > num3 + 100:  # 确保结果为正数
                            answer = num1 + num2 - num3
                            problem = f"{num1} {op1} {num2} {op2} {num3} "
                            return problem, answer, difficulty
                    elif op1 == '-' and op2 == '+':
                        if num1 > num2:  # 确保第一次减法结果为正数
                            answer = num1 - num2 + num3
                            problem = f"{num1} {op1} {num2} {op2} {num3} "
                            return problem, answer, difficulty
                    else:  # op1 == '-' and op2 == '-'
                        if num1 > num2 + num3 + 100:  # 确保结果为正数
                            answer = num1 - num2 - num3
                            problem = f"{num1} {op1} {num2} {op2} {num3} "
                            return problem, answer, difficulty

                # 如果当前组合不成功，重新选择运算符
                op1 = random.choice(operations)
                op2 = random.choice(operations)

        # 如果多次尝试都失败，生成一个简单的混合运算
        num1 = random.randint(1000, 9999)
        num2 = random.randint(100, 999)
        num3 = random.randint(1000, 9999)
        answer = num1 + num2 * 10  # 确保是正数
        problem = f"{num1} + {num2} × 10 + {num3} "
        return problem, answer, difficulty

    def generate_problems(self, count=30, difficulty="random", auto_adjust_count=False):
        """生成指定数量和难度的题目

        Args:
            count (int): 题目数量
            difficulty (str or list): 难度级别
            auto_adjust_count (bool): 是否自动调整题目数量以优化分页
        """
        if auto_adjust_count:
            first_page_count = 52
            # (616-52)/10 = 56.4
            standard_page_count = 56
            # min_fill_ratio = 0.98
            min_fill_ratio = 1

            if count > first_page_count:
                min_last_page_count = math.ceil(standard_page_count * min_fill_ratio)

                remaining_problems = count - first_page_count
                last_page_problems = remaining_problems % math.floor(standard_page_count)

                if 0 < last_page_problems < min_last_page_count:
                    problems_to_add = min_last_page_count - last_page_problems
                    adjusted_count = count + problems_to_add
                    print(f"提示：原题目数 {count} 会导致最后一页填充不足。")
                    print(f"      为优化分页，已自动调整为 {adjusted_count} 题。")
                    count = adjusted_count

        self.problems = []
        self.answers = []
        self.selected_difficulty = difficulty if isinstance(difficulty, list) else [difficulty]

        operations = [
            self.generate_addition,
            self.generate_subtraction,
            self.generate_multiplication,
            self.generate_division,
            self.generate_mixed_operation,
        ]

        # 处理难度参数
        if difficulty == "random":
            selected_difficulties = ["easy", "medium", "hard", "harder"]
        elif isinstance(difficulty, str):
            selected_difficulties = [difficulty]
        elif isinstance(difficulty, list):
            selected_difficulties = difficulty
        else:
            selected_difficulties = ["easy"]  # 默认

        # 验证难度级别有效性
        valid_difficulties = ["easy", "medium", "hard", "harder"]
        selected_difficulties = [d for d in selected_difficulties if d in valid_difficulties]
        if not selected_difficulties:
            selected_difficulties = ["easy"]  # 如果没有有效难度，使用默认

        print(f"选择的难度级别: {selected_difficulties}")
        print(f"题目数量: {count}")

        for _ in range(count):
            # 从选择的难度列表中随机选择
            selected_difficulty = random.choice(selected_difficulties)

            # 对于harder难度，优先使用混合运算
            if selected_difficulty == "harder" and random.random() < 0.7:
                operation = self.generate_mixed_operation
            else:
                operation = random.choice(operations[:4])  # 只使用基本运算

            problem, answer, diff_level = operation(selected_difficulty)

            # 如果生成失败，使用加法作为备选
            if problem is None:
                problem, answer, diff_level = self.generate_addition(selected_difficulty)

            self.problems.append((problem, answer, diff_level))
            self.answers.append(answer)

    def create_columns(self, doc, num_columns=2):
        """创建多栏布局"""
        # 添加分节符
        doc.add_section(WD_SECTION.CONTINUOUS)

        # 获取当前节
        section = doc.sections[-1]

        # 设置多栏布局
        sect_pr = section._sectPr
        if sect_pr is None:
            return

        # 创建栏属性
        from docx.oxml import parse_xml
        from docx.oxml.shared import qn

        # 根据栏数调整间距，多于两栏时使用较小间距
        space = "360" if num_columns > 2 else "720"

        # 添加多栏布局
        cols_xml = parse_xml(
            f'<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:num="{num_columns}" w:space="{space}"/>'
        )
        sect_pr.append(cols_xml)

    def add_header_footer(self, doc, is_answer=False, difficulty=None):
        """添加页眉和页脚"""
        # --- 添加页眉 ---
        header = doc.sections[0].header
        # 清除默认段落，为表格做准备
        header.paragraphs[0].clear()

        # 最终解决方案：使用一个1x2的隐形表格来实现精确的左右对齐
        section = doc.sections[0]
        table_width = section.page_width - section.left_margin - section.right_margin
        table = header.add_table(rows=1, cols=2, width=table_width)

        # 获取左右两个单元格
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        # 关键优化：将单元格的内部边距（padding）设置为0，以最大程度减小表格高度
        for cell in [left_cell, right_cell]:
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'bottom']:
                mar_el = OxmlElement(f'w:{margin}')
                mar_el.set(qn('w:w'), '0')
                mar_el.set(qn('w:type'), 'dxa')
                tcMar.append(mar_el)
            tcPr.append(tcMar)

        # 配置左侧单元格，并移除段落间距以减小高度
        left_paragraph = left_cell.paragraphs[0]
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_format_left = left_paragraph.paragraph_format
        p_format_left.space_before = Pt(0)
        p_format_left.space_after = Pt(0)
        left_run = left_paragraph.add_run("Date: ")
        left_run.font.size = Pt(9)

        # 配置右侧单元格，并移除段落间距以减小高度
        right_paragraph = right_cell.paragraphs[0]
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_format_right = right_paragraph.paragraph_format
        p_format_right.space_before = Pt(0)
        p_format_right.space_after = Pt(0)
        right_text = "Large Numbers Math Problems" if not is_answer else "Large Numbers Math Answers"

        # --- 添加难度等级 ---
        if difficulty:
            difficulty_map = {"easy": "*", "medium": "**", "hard": "***", "harder": "****"}
            difficulty_order = ["easy", "medium", "hard", "harder"]

            actual_difficulties = set(difficulty)
            if "random" in actual_difficulties:
                actual_difficulties.update(difficulty_order)

            present_difficulties = [d for d in difficulty_order if d in actual_difficulties]
            if present_difficulties:
                star_strings = [difficulty_map.get(d, "") for d in present_difficulties]
                stars = " ".join(star_strings)
                right_text += f" {stars}"

        right_run = right_paragraph.add_run(right_text)
        right_run.font.size = Pt(9)

        # 移除表格边框，使其不可见
        # 通过直接操作OXML来设置所有边框为'nil'
        tbl = table._tbl
        # 兼容不同版本的 python-docx：先检查 tblPr 是否存在，不存在则添加
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = tbl._add_tblPr()

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'nil')
            tblBorders.append(border_el)
        tblPr.append(tblBorders)

        # --- 添加页脚 ---
        footer = doc.sections[0].footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.clear()

        # 添加生成时间
        time_text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        time_run = footer_paragraph.add_run(time_text)
        time_run.font.size = Pt(9)

        # 添加页号 (页号/总页数)
        footer_paragraph.add_run("  |  ")

        # 创建一个内部函数来添加域代码，避免重复
        def add_field(run, field_code):
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = f' {field_code} '
            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            run._r.extend([fldChar_begin, instrText, fldChar_end])

        # 添加 "当前页号"
        page_num_run = footer_paragraph.add_run()
        page_num_run.font.size = Pt(9)
        add_field(page_num_run, "PAGE")

        # 添加 "/"
        slash_run = footer_paragraph.add_run("/")
        slash_run.font.size = Pt(9)

        # 添加 "总页数"
        total_pages_run = footer_paragraph.add_run()
        total_pages_run.font.size = Pt(9)
        add_field(total_pages_run, "NUMPAGES")

        # 在题目文档的页脚添加难度说明
        if not is_answer:
            footer_paragraph.add_run("  |  ")
            explanation_run = footer_paragraph.add_run("难度说明: ")
            explanation_run.font.size = Pt(9)

            for diff, color in self.difficulty_colors.items():
                example_run = footer_paragraph.add_run("= ")
                example_run.font.color.rgb = color
                example_run.bold = True
                name_run = footer_paragraph.add_run(f"{self.difficulty_names[diff]}  ")
                name_run.font.size = Pt(9)

        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def format_number_with_dot_alignment(self, number, total_digits=2):
        """格式化数字以实现小数点对齐"""
        return f"{number:>{total_digits}d}."

    def add_problem_with_gray_equal(self, paragraph, problem_text, difficulty, problem_number, total_problems):
        """添加带有灰度等号的题目（小数点对齐）"""
        # 计算需要的数字位数
        total_digits = len(str(total_problems))

        # 添加格式化后的题号（小数点对齐）
        formatted_number = self.format_number_with_dot_alignment(problem_number, total_digits)
        run = paragraph.add_run(formatted_number)
        run.bold = True
        run.font.color.rgb = RGBColor(170, 170, 170)  # 设置题号为浅灰色

        # 添加题目和灰度等号
        run = paragraph.add_run(f"  {problem_text}")
        equal_run = paragraph.add_run("= ")
        equal_run.font.color.rgb = self.difficulty_colors[difficulty]
        equal_run.bold = True

    def _setup_document(self, title_text, is_answer=False, num_columns=2, difficulty=None):
        """初始化Word文档，包含标准格式设置"""
        doc = Document()
        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # 添加页眉和页脚
        self.add_header_footer(doc, is_answer=is_answer, difficulty=difficulty)

        # 设置标题
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # doc.add_paragraph()  # 空行

        # 创建多栏布局
        self.create_columns(doc, num_columns)
        return doc

    def _populate_problems_document(self, doc):
        """向文档中填充所有题目"""
        total_problems = len(self.problems)
        for i, (problem, answer, difficulty) in enumerate(self.problems, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            self.add_problem_with_gray_equal(p, problem, difficulty, i, total_problems)

            # 每10题添加一个较大的空行作为分组分隔
            if i % 10 == 0 and i < total_problems:
                separator = doc.add_paragraph()
                separator.paragraph_format.space_after = Pt(9)

    def _populate_answers_document(self, doc):
        """向文档中填充所有答案"""
        total_problems = len(self.problems)
        answer_font_name = 'Calibri'  # 为所有答案指定统一字体

        for i, (answer, (problem, _, difficulty)) in enumerate(zip(self.answers, self.problems), 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)

            # 添加格式化后的题号
            formatted_number = self.format_number_with_dot_alignment(i, len(str(total_problems)))
            run = p.add_run(formatted_number)
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.name = answer_font_name

            # 添加答案
            p.add_run("  ")  # 添加一些空格
            answer_run = p.add_run(str(answer))
            answer_run.font.color.rgb = self.difficulty_colors[difficulty]
            answer_run.bold = True
            answer_run.font.size = Pt(10.5)
            answer_run.font.name = answer_font_name

            # 每10题添加一个较大的空行作为分组分隔
            if i % 10 == 0 and i < total_problems:
                separator = doc.add_paragraph()
                separator.paragraph_format.space_after = Pt(9)

    def _populate_problems_and_answers_document(self, doc):
        """向文档中填充所有题目，并在新页面上以四栏格式附加答案"""
        # 1. 填充所有题目 (在默认的双栏布局中)
        self._populate_problems_document(doc)

        # 2. 添加一个分页分节符，以便为答案设置新的布局
        doc.add_section(WD_SECTION.NEW_PAGE)

        # 3. 为新的答案部分设置四栏布局
        section = doc.sections[-1]
        sect_pr = section._sectPr
        from docx.oxml import parse_xml

        space = "360"  # 4栏使用较小的间距
        cols_xml = parse_xml(
            f'<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:num="4" w:space="{space}"/>'
        )
        sect_pr.append(cols_xml)

        # 4. 添加答案部分的标题
        answer_title = doc.add_heading('答案', level=1)
        answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 5. 填充所有答案 (将在新的四栏布局中)
        self._populate_answers_document(doc)

    def save_to_word(self, filepath=None):
        """将题目和答案保存到Word文档（两栏布局，英文命名）

        Args:
            filepath (str, optional): 文件保存路径。如果为None，则保存到系统临时目录
        """
        timestamp = datetime.now().strftime("%Y%m%d")

        # 确定保存目录
        save_dir = filepath or tempfile.gettempdir()
        if not path.exists(save_dir):
            os.makedirs(save_dir)

        # 生成文件名
        problems_filename = path.join(save_dir, f"Large_Numbers_Math_Problems_{timestamp}.docx")
        answers_filename = path.join(save_dir, f"Large_Numbers_Math_Answers_{timestamp}.docx")
        problems_and_answers_filename = path.join(save_dir, f"Large_Numbers_Math_Problems_and_Answers_{timestamp}.docx")

        # 创建并保存题目文档 (2栏)
        doc_problems = self._setup_document(
            '大数计算', is_answer=False, num_columns=2, difficulty=self.selected_difficulty
        )
        self._populate_problems_document(doc_problems)
        doc_problems.save(problems_filename)

        # 创建并保存答案文档 (4栏)
        doc_answers = self._setup_document(
            '大数计算答案', is_answer=True, num_columns=4, difficulty=self.selected_difficulty
        )
        self._populate_answers_document(doc_answers)
        doc_answers.save(answers_filename)

        # 创建并保存题目与答案合并文档 (2栏)
        doc_problems_and_answers = self._setup_document(
            '大数计算 (含答案)', is_answer=False, num_columns=2, difficulty=self.selected_difficulty
        )
        self._populate_problems_and_answers_document(doc_problems_and_answers)
        doc_problems_and_answers.save(problems_and_answers_filename)

        print(f"Problems saved to: {problems_filename}")
        print(f"Answers saved to: {answers_filename}")
        print(f"Problems and Answers saved to: {problems_and_answers_filename}")
        print(f"题目已保存到: {problems_filename}")
        print(f"答案已保存到: {answers_filename}")
        print(f"题目和答案已保存到: {problems_and_answers_filename}")

        return problems_filename, answers_filename, problems_and_answers_filename


def parse_arguments():
    """解析命令行参数，支持难度等级的数字别名。"""
    # --- 命令行参数预处理 ---
    difficulty_map = {
        "1": "easy",
        "2": "medium",
        "3": "hard",
        "4": "harder",
        "5": "random",
    }
    # 遍历sys.argv[1:]，跳过脚本名称本身
    for i in range(1, len(sys.argv)):
        if sys.argv[i] in difficulty_map:
            sys.argv[i] = difficulty_map[sys.argv[i]]
    # --- 预处理结束 ---

    parser = argparse.ArgumentParser(
        description="生成大数计算题库Word文档。", formatter_class=argparse.RawTextHelpFormatter  # 保持帮助信息格式
    )
    parser.add_argument(
        "-c",
        "--count",
        type=int,
        default=660,
        help="要生成的题目数量 (默认: 660)",
    )
    parser.add_argument(
        "-d",
        "--difficulty",
        nargs="+",
        # default=["hard", "medium"],
        default=["easy", "medium"],
        choices=["easy", "medium", "hard", "harder", "random"],
        help='''难度级别，可多选 (默认: hard medium)。
可选: "easy", "medium", "hard", "harder", "random"。
支持数字别名: 1=easy, 2=medium, 3=hard, 4=harder, 5=random''',
    )
    parser.add_argument(
        "--no-adjust",
        action="store_true",
        help="禁用自动调整题目数量以优化分页的功能",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=str,
        default=None,
        help="输出Word文档的目录路径 (默认: 系统临时目录)",
    )

    return parser.parse_args()


# 使用示例
if __name__ == "__main__":
    args = parse_arguments()

    generator = LargeNumberMathGenerator()

    # 根据命令行参数生成题目
    print("--- 开始生成题目 ---")
    generator.generate_problems(
        count=args.count,
        difficulty=args.difficulty,
        auto_adjust_count=not args.no_adjust,
    )

    # 保存文件
    problems_file, answers_file, problems_and_answers_file = generator.save_to_word(filepath=args.output)

    print("\n--- 生成完毕 ---")
    print(f"题目文档已保存至: {problems_file}")
    print(f"答案文档已保存至: {answers_file}")
    print(f"题目和答案文档已保存至: {problems_and_answers_file}")

    print("\n文档特点:")
    print("- 页眉: Large Numbers Math Problems / Large Numbers Math Answers")
    print("- 页脚: 英文格式的生成时间")
    print("- 题目序号小数点对齐")
    print("- 等号使用不同灰度表示难度")

    print("\n难度说明:")
    print("- 题目文档: 彩色'='表示难度")
    print("- 答案文档: 彩色答案表示难度")
    print("- 颜色对应: 浅灰色=简单, 中灰色=中等, 深灰色=困难, 黑色=更难（四则混合）")
    print("黑色等号   = 更难题目（四则混合运算）")
