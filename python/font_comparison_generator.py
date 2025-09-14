
from docx import Document
from docx.shared import Pt
import os
import tempfile

def create_font_comparison_doc():
    """
    Creates a Word document to showcase different fonts for math problems.
    """
    doc = Document()
    doc.add_heading('字体样式对比', 0)
    p = doc.add_paragraph('以下为几种不同字体的数学题目样式，以供您参考选择。')
    p.add_run('（建议将Word窗口放大查看）').italic = True


    fonts_to_compare = ['Calibri', 'Times New Roman', 'Arial', 'Consolas', 'Cambria']
    sample_problems = [
        '1.  8192 + 4327 =',
        '2.  9017 - 5823 =',
        '3.  485 × 62 =',
        '4.  7254 ÷ 9 =',
        '5.  1l0O 与 1100 (易混淆字符对比)'
    ]

    for font_name in fonts_to_compare:
        try:
            doc.add_heading(f'字体: {font_name}', level=2)
            for problem in sample_problems:
                p = doc.add_paragraph()
                run = p.add_run(problem)
                font = run.font
                font.name = font_name
                font.size = Pt(12) # 使用稍大字号以便对比
        except Exception:
            # If a font is not installed, Word will use a fallback.
            doc.add_paragraph(f"(注意: 无法应用字体 '{font_name}'。您的系统中可能未安装此字体。)")


    # Save the document to a temporary file
    filepath = os.path.join(tempfile.gettempdir(), "Font_Comparison_Example.docx")
    doc.save(filepath)
    return filepath

if __name__ == '__main__':
    doc_path = create_font_comparison_doc()
    print(f"Comparison document saved to: {doc_path}")
